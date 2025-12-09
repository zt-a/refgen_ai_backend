from docx import Document
from matplotlib import font_manager as fm
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from src.refprint.LatexConverter.latex2omml import LatexConverter
import matplotlib.pyplot as plt


class RefPrint:
    latex_converter = LatexConverter()

    def __init__(self, doc: Document):
        self.doc = doc
        self.contents = []

        self.section = self.doc.sections[0]

        # Настройка шрифта для формул matplotlib
        times_new_roman = './Times New Roman Regular.ttf'
        fm.FontProperties(fname=times_new_roman)
        plt.rcParams['mathtext.fontset'] = 'custom'
        plt.rcParams['mathtext.rm'] = 'Times New Roman'
        plt.rcParams['mathtext.it'] = 'Times New Roman:italic'
        plt.rcParams['mathtext.bf'] = 'Times New Roman:bold'

        # Поля документа
        self.section.left_margin = Cm(3)
        self.section.right_margin = Cm(1)
        self.section.top_margin = Cm(2)
        self.section.bottom_margin = Cm(2)

    @staticmethod
    def set_run_font(run, size=14, name='Times New Roman', bold=False, italic=False):
        run.font.name = name
        run.font.size = Pt(max(size, 14))
        run.bold = bold
        run.italic = italic
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), name)

    @staticmethod
    def set_paragraph_alignment(para):
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def ref_add(self, data: str):
        """Добавляем HTML-контент"""
        if data and data.strip():
            self.contents.append(data.strip())

    def add_footer(self):
        self.section.different_first_page_header_footer = True
        footer = self.section.footer
        para_left = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para_left.text = "refgen"
        para_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        para_right = footer.add_paragraph()
        para_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = para_right.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2])

    def _process_inline_tags(self, element):
        """Рекурсивная обработка <b> и <i> внутри <p> или <li>"""
        runs = []
        if element.name is None:
            runs.append((element.string or "", False, False))
        else:
            bold = element.name == 'b'
            italic = element.name == 'i'
            for child in element.children:
                child_runs = self._process_inline_tags(child)
                for text, b, i in child_runs:
                    runs.append((text, bold or b, italic or i))
        return runs

    def ref_print(self) -> Document:
        # Стиль документа
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1

        for html_data in self.contents:
            soup = BeautifulSoup(html_data, 'html.parser')
            content_tag = soup.find('content')
            formulas_tag = soup.find('formulas')

            if not content_tag:
                continue

            # Формулы
            formulas_dict = {}
            if formulas_tag:
                for latex_tag in formulas_tag.find_all('latex'):
                    formulas_dict[latex_tag['id']] = latex_tag.get_text(strip=True)

            for element in content_tag.children:
                # Игнорируем пустые текстовые узлы
                if element.name is None:
                    text = (element.string or '').strip()
                    if not text:
                        continue
                    para = self.doc.add_paragraph(text)
                    self.set_paragraph_alignment(para)
                    continue

                text = element.get_text(strip=True)
                if not text and element.name not in ('br', 'table', 'ul', 'ol', 'formula'):
                    continue

                # Заголовки
                if element.name in ('h1','h2','h3'):
                    level_map = {'h1': 1, 'h2': 2, 'h3': 3}
                    size_map = {'h1': 20, 'h2': 18, 'h3': 16}
                    para = self.doc.add_heading(level=level_map[element.name])
                    run = para.add_run(text)
                    self.set_run_font(run, size=size_map[element.name])
                    self.set_paragraph_alignment(para)

                # Параграфы
                elif element.name == 'p':
                    para_text = ''.join([t for t, _, _ in self._process_inline_tags(element)]).strip()
                    if not para_text:
                        continue
                    para = self.doc.add_paragraph()
                    for run_text, bold, italic in self._process_inline_tags(element):
                        if run_text.strip():
                            run = para.add_run(run_text)
                            self.set_run_font(run, size=14, bold=bold, italic=italic)
                    self.set_paragraph_alignment(para)

                # Списки
                elif element.name in ('ul','ol'):
                    style_name = 'List Bullet' if element.name == 'ul' else 'List Number'
                    for li in element.find_all('li', recursive=False):
                        li_text = ''.join([t for t, _, _ in self._process_inline_tags(li)]).strip()
                        if not li_text:
                            continue
                        para = self.doc.add_paragraph(style=style_name)
                        for run_text, bold, italic in self._process_inline_tags(li):
                            if run_text.strip():
                                run = para.add_run(run_text)
                                self.set_run_font(run, size=14, bold=bold, italic=italic)

                # Формулы
                elif element.name == 'formula' and formulas_dict:
                    formula_id = element.get('id')
                    latex = formulas_dict.get(formula_id)
                    if latex:
                        para = self.doc.add_paragraph()
                        omml_xml = self.latex_converter.convert(latex)
                        para._p.append(omml_xml)

                # Таблицы
                elif element.name == 'table':
                    rows = element.find_all('tr', recursive=False)
                    if not rows:
                        continue
                    table = self.doc.add_table(rows=0, cols=len(rows[0].find_all(['th','td'])))
                    for row in rows:
                        cells = row.find_all(['th','td'])
                        table_row = table.add_row().cells
                        for i, cell in enumerate(cells):
                            para = table_row[i].paragraphs[0]
                            for run_text, bold, italic in self._process_inline_tags(cell):
                                if run_text.strip():
                                    run = para.add_run(run_text)
                                    self.set_run_font(run, size=14, bold=bold, italic=italic)

                # Разрыв страницы
                elif element.name == 'br':
                    para = self.doc.add_paragraph()
                    para.add_run().add_break(WD_BREAK.PAGE)

        return self.doc

