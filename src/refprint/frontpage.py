from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


class FrontPage:
    def __init__(
        self, 
        university: str | None,
        faculty: str | None,
        subject: str | None,
        topic: str | None,
        course: str | None,
        performed_by: str | None,
        checked_by: str | None,
        group: str | None,
        city: str | None,
        year: str | None,
        doc: Document
    ):
        self.university = university
        self.faculty = faculty
        self.subject = subject
        self.topic = topic
        self.course = course
        self.performed_by = performed_by
        self.checked_by = checked_by
        self.group = group
        self.city = city
        self.year = year
        self.doc: Document = doc

    def write(self):
        # Настройка шрифта документа
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Безопасные значения
        university = self.university or "_________________________________________________"
        subject = self.subject or "_______________"
        faculty = self.faculty or "__________________________"
        topic = self.topic or "_________________________________________"
        performed_by = self.performed_by or "_________________________"
        group = self.group or "_____________________________"
        checked_by = self.checked_by or "_________________________"
        city = self.city or "Бишкек"
        year = self.year or "2025"

        # ----------------------
        # Верхний центр
        self.doc.add_paragraph("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КР").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("Государственное учреждение высшего профессионального образования").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(university).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем вертикальный отступ (3 пустых параграфа)
        for _ in range(10):
            self.doc.add_paragraph()

        # ----------------------
        # Центр — РЕФЕРАТ
        p = self.doc.add_paragraph("РЕФЕРАТ")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(80)
        run.font.bold = True

        # Данные
        self.doc.add_paragraph(f"Дисциплина: {subject}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"Кафедра: {faculty}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph(f"Тема: {topic}").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые параграфы для вертикального смещения
        for _ in range(10):
            self.doc.add_paragraph()

        # ----------------------
        # Правая колонка — ближе к низу
        p1 = self.doc.add_paragraph(f"Выполнил: {performed_by}")
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2 = self.doc.add_paragraph(f"Группа: {group}")
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3 = self.doc.add_paragraph(f"Проверил: {checked_by}")
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Дополнительные пустые параграфы для смещения города и года вниз
        for _ in range(3):
            self.doc.add_paragraph()

        # ----------------------
        # Нижний центр — город и год
        p = self.doc.add_paragraph(f"{city}, {year}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

        # Разрыв страницы
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        return self.doc
